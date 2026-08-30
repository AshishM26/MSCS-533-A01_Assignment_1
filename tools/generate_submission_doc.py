#!/usr/bin/env python3
"""Generate the local Blackboard Word submission from repository contents."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as error:
    raise SystemExit(
        "python-docx is required. Install it with: "
        "python3 -m pip install python-docx"
    ) from error


COURSE = "2026 Fall - Software Engineering and Multiplatform App Development"
COURSE_CODE = "MSCS-533-A01"
ASSIGNMENT = "Hands-on Assignment 1: Construct Your First Flutter App using Dart"
STUDENT = "Ashish Mahajan"
OUTPUT_PATH = Path("submission/MSCS533_HandsOn_Assignment1_AM.docx")
SCREENSHOT_PATH = Path("screenshots/measures_converter.png")
MANIFEST_PATH = Path("android/app/src/main/AndroidManifest.xml")
SOURCE_GLOB = "lib/**/*.dart"

BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = "F2F4F7"
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def repository_root() -> Path:
    root = Path(__file__).resolve().parents[1]
    if not (root / ".git").exists():
        raise SystemExit(f"Repository metadata was not found at: {root}")
    return root


def run_git(root: Path, *arguments: str) -> str:
    result = subprocess.run(
        ["git", *arguments],
        cwd=root,
        check=False,
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        message = result.stderr.strip() or result.stdout.strip()
        raise SystemExit(f"Git command failed: git {' '.join(arguments)}\n{message}")
    return result.stdout.strip()


def browser_repository_url(remote_url: str) -> str:
    remote_url = remote_url.strip()
    ssh_match = re.fullmatch(r"git@github\.com:(.+?)(?:\.git)?", remote_url)
    if ssh_match:
        return f"https://github.com/{ssh_match.group(1)}"

    ssh_url_match = re.fullmatch(
        r"ssh://git@github\.com/(.+?)(?:\.git)?", remote_url
    )
    if ssh_url_match:
        return f"https://github.com/{ssh_url_match.group(1)}"

    if remote_url.startswith("https://github.com/"):
        return remote_url.removesuffix(".git")

    raise SystemExit(
        "The origin remote is not a supported GitHub URL: " f"{remote_url}"
    )


def require_inputs(root: Path) -> list[Path]:
    required_paths = [root / SCREENSHOT_PATH, root / MANIFEST_PATH]
    missing = [path for path in required_paths if not path.is_file()]
    if missing:
        relative = ", ".join(str(path.relative_to(root)) for path in missing)
        raise SystemExit(f"Required input file(s) missing: {relative}")

    source_files = sorted(root.glob(SOURCE_GLOB))
    if not source_files:
        raise SystemExit(f"No Dart source files matched {SOURCE_GLOB}")
    return source_files


def set_run_font(
    run,
    *,
    name: str,
    size: float,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor = BLACK,
) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (15, 16, 8),
        "Heading 2": (12, 12, 5),
        "Heading 3": (12, 8, 4),
    }
    for style_name, (size, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = styles.add_style("Repository Code", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(8)
    code_style.font.color.rgb = BLACK
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    code_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code_style.paragraph_format.keep_together = False
    code_style.paragraph_format.widow_control = False
    word_wrap = OxmlElement("w:wordWrap")
    word_wrap.set(qn("w:val"), "0")
    code_style._element.get_or_add_pPr().append(word_wrap)

    caption_style = styles.add_style("Academic Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = "Times New Roman"
    caption_style.font.size = Pt(10)
    caption_style.font.italic = True
    caption_style.font.color.rgb = BLACK
    caption_style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption_style.paragraph_format.space_before = Pt(4)
    caption_style.paragraph_format.space_after = Pt(8)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run(
        "MSCS-533-A01\nSoftware Engineering and Multiplatform App Development"
    )
    set_run_font(
        header_run, name="Times New Roman", size=9, color=MUTED
    )

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    footer_run = footer_paragraph.add_run(
        "Ashish Mahajan\nUniversity of the Cumberlands\nPage "
    )
    set_run_font(footer_run, name="Times New Roman", size=9, color=MUTED)
    add_page_number_field(footer_paragraph)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, name="Times New Roman", size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    visible = OxmlElement("w:t")
    visible.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, visible, end])


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    run_properties.extend([fonts, color, underline, size])
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_label_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, name="Times New Roman", size=12, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, name="Times New Roman", size=12)


def add_title_block(
    document: Document, repository_url: str, commit_hash: str
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(COURSE)
    set_run_font(run, name="Times New Roman", size=14, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(f"({COURSE_CODE})")
    set_run_font(run, name="Times New Roman", size=13)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("Hands-on Assignment 1:")
    set_run_font(run, name="Times New Roman", size=16, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(22)
    run = paragraph.add_run("Construct Your First Flutter App using Dart")
    set_run_font(run, name="Times New Roman", size=16, bold=True)

    add_label_value(document, "Name", STUDENT)
    add_label_value(
        document,
        "University",
        "University of the Cumberlands",
    )
    add_label_value(document, "Course", COURSE_CODE)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run("GitHub Assignment: ")
    set_run_font(label_run, name="Times New Roman", size=12, bold=True)
    add_hyperlink(paragraph, repository_url, repository_url)
    add_label_value(document, "Final Commit", commit_hash)


def add_section_heading(
    document: Document,
    number: int,
    title: str,
    *,
    page_break_before: bool = False,
) -> None:
    heading = document.add_heading(f"{number}. {title}", level=1)
    heading.paragraph_format.page_break_before = page_break_before


def set_image_alt_text(inline_shape, description: str) -> None:
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", "Measures Converter application output")


def add_source_code(document: Document, root: Path, source_files: list[Path]) -> None:
    for source_file in source_files:
        relative_path = source_file.relative_to(root).as_posix()
        document.add_heading(relative_path, level=2)
        contents = source_file.read_text(encoding="utf-8")
        for line in contents.splitlines():
            paragraph = document.add_paragraph(style="Repository Code")
            paragraph.add_run(line)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr

    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths_dxa)))

    table_indent = table_properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    cell_margins = table_properties.find(qn("w:tblCellMar"))
    if cell_margins is None:
        cell_margins = OxmlElement("w:tblCellMar")
        table_properties.append(cell_margins)
    for margin_name, margin_value in CELL_MARGIN_DXA.items():
        margin = cell_margins.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            cell_margins.append(margin)
        margin.set(qn("w:w"), str(margin_value))
        margin.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for grid_column in list(grid):
        grid.remove(grid_column)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths_dxa[index]))


def format_table_text(table) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        name="Times New Roman",
                        size=10,
                        bold=row_index == 0,
                    )
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)


def add_validation_table(document: Document) -> None:
    rows = [
        (
            "Formatting",
            "dart format .",
            "PASS - Formatted 5 files; 0 changed.",
        ),
        (
            "Static analysis",
            "flutter analyze",
            "PASS - No issues found.",
        ),
        (
            "Automated tests",
            "flutter test",
            "PASS - All 11 unit and widget tests passed.",
        ),
        (
            "Runtime",
            "flutter run -d emulator-5554 --no-resident",
            "PASS - Built, installed, and launched on Android 17 (API 37); no "
            "Flutter exception or overflow appeared in the final log scan.",
        ),
        (
            "Length conversion",
            "100 meters -> feet",
            "PASS - 328.084 feet (live emulator and automated test).",
        ),
        (
            "Length conversion",
            "1 mile -> kilometers",
            "PASS - 1.609344 kilometers (automated test).",
        ),
        (
            "Mass conversion",
            "1 kilogram -> pounds",
            "PASS - 2.204623 pounds (live emulator and automated test).",
        ),
        (
            "Mass conversion",
            "1 pound -> kilograms",
            "PASS - 0.45359237 kilograms (automated test).",
        ),
        (
            "Same-unit conversion",
            "42.5 ounces -> ounces",
            "PASS - Returned 42.5 ounces (automated test).",
        ),
        (
            "Input validation",
            "Empty and invalid input",
            "PASS - Displayed 'Enter a numeric value.' (live emulator and widget test).",
        ),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Validation", "Command / Check", "Result"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        set_cell_shading(table.rows[0].cells[index], LIGHT_GRAY)

    for validation, check, result in rows:
        cells = table.add_row().cells
        cells[0].text = validation
        cells[1].text = check
        cells[2].text = result

    set_table_geometry(table, [1900, 3000, 4460])
    format_table_text(table)


def build_document(
    root: Path,
    source_files: list[Path],
    repository_url: str,
    branch: str,
    commit_hash: str,
) -> Document:
    document = Document()
    configure_styles(document)
    configure_page(document)

    properties = document.core_properties
    properties.title = ASSIGNMENT
    properties.subject = "MSCS-533-A01 Hands-on Assignment 1 submission"
    properties.author = STUDENT
    properties.keywords = "Flutter, Dart, Measures Converter, MSCS-533-A01"
    properties.comments = "Generated from the final Git repository contents."

    add_title_block(document, repository_url, commit_hash)

    add_section_heading(document, 1, "Application Overview")
    document.add_paragraph(
        "Measures Converter is a Flutter application written in Dart that converts "
        "metric and imperial measurements. It supports length and mass units and "
        "keeps destination choices within the source unit's category, preventing "
        "invalid length-to-mass conversions."
    )
    document.add_paragraph(
        "The interface accepts numeric and decimal input, provides From and To unit "
        "selectors, performs conversions on demand, and displays a clearly formatted "
        "result. Same-unit and repeated conversions are also supported."
    )

    add_section_heading(document, 2, "Application Output")
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    image_run = image_paragraph.add_run()
    image = image_run.add_picture(str(root / SCREENSHOT_PATH), width=Inches(2.55))
    set_image_alt_text(
        image,
        "Measures Converter on an Android emulator showing 115 kilometers converted "
        "to 71.457687 miles.",
    )
    caption = document.add_paragraph(style="Academic Caption")
    caption.add_run("Figure 1. Measures Converter running on the Android emulator.")

    add_section_heading(document, 3, "Implementation Design")
    document.add_paragraph(
        "The declarative Flutter interface is implemented as a StatefulWidget that "
        "owns the numeric input controller, selected units, validation message, and "
        "displayed result. From and To dropdowns update this state, while the "
        "destination list is derived from the selected unit category so length and "
        "mass units cannot be mixed."
    )
    document.add_paragraph(
        "ConversionUnit defines each unit and its factor relative to a category base "
        "unit: meters for length and kilograms for mass. ConversionService validates "
        "compatibility, converts the source value to the base unit, and then converts "
        "to the destination. This keeps conversion formulas separate from the UI and "
        "avoids duplicated pair-by-pair calculations."
    )

    add_section_heading(document, 4, "Dart Source Code", page_break_before=True)
    document.add_paragraph(
        "The following source is read directly from the final repository so the "
        "submission matches the committed implementation."
    )
    add_source_code(document, root, source_files)

    add_section_heading(document, 5, "Android Manifest", page_break_before=True)
    document.add_paragraph(
        "Complete contents of android/app/src/main/AndroidManifest.xml:"
    )
    for line in (root / MANIFEST_PATH).read_text(encoding="utf-8").splitlines():
        paragraph = document.add_paragraph(style="Repository Code")
        paragraph.add_run(line)

    add_section_heading(document, 6, "Testing and Validation", page_break_before=True)
    document.add_paragraph(
        "Final validation was completed on August 30, 2026. The results below are "
        "from the final audit of this repository."
    )
    add_validation_table(document)
    document.add_paragraph(
        "Representative conversion checks: 100 meters to 328.084 feet; 1 mile to "
        "1.609344 kilometers; 1 kilogram to approximately 2.204623 pounds; 1 pound "
        "to 0.45359237 kilograms; 10 feet to 3.048 meters; and 100 grams to "
        "approximately 3.527396 ounces. A same-unit conversion of 42.5 ounces also "
        "returned the original value. Floating-point comparisons used tolerances."
    )

    add_section_heading(document, 7, "Coding Practices")
    document.add_paragraph(
        "The implementation separates Flutter UI and interaction state from the "
        "conversion model and conversion service. Classes, methods, and variables use "
        "meaningful Effective Dart naming and null-safe types. Immutable unit "
        "definitions and const constructors are used where appropriate."
    )
    document.add_paragraph(
        "Each conversion first maps the source value to a category base unit (meter "
        "for length or kilogram for mass) and then maps that value to the destination "
        "unit. This avoids duplicated pair-specific formulas. Input validation rejects "
        "empty, nonnumeric, and non-finite values; incompatible categories are also "
        "rejected. Concise comments document intent, while formatting, static analysis, "
        "unit tests, and widget tests provide automated quality checks."
    )

    add_section_heading(document, 8, "GitHub Repository")
    paragraph = document.add_paragraph()
    paragraph.add_run("Repository URL: ").bold = True
    add_hyperlink(paragraph, repository_url, repository_url)
    add_label_value(document, "Branch", branch)
    add_label_value(document, "Final commit", commit_hash)

    add_section_heading(document, 9, "Development Tools and AI Assistance")
    document.add_paragraph(
        "Development used Visual Studio Code, Flutter, Dart, Android Studio and the "
        "Android Emulator, Git, GitHub, and AI-assisted development tools for setup "
        "guidance, implementation support, review, and learning. The final "
        "implementation was reviewed, tested, and understood by the student."
    )

    return document


def repository_code_lines(document: Document, heading_text: str) -> list[str]:
    paragraphs = document.paragraphs
    heading_index = next(
        (index for index, paragraph in enumerate(paragraphs) if paragraph.text == heading_text),
        None,
    )
    if heading_index is None:
        raise SystemExit(f"Document validation failed: missing heading {heading_text}")

    code_lines: list[str] = []
    code_started = False
    for paragraph in paragraphs[heading_index + 1 :]:
        if paragraph.style.name == "Repository Code":
            code_started = True
            code_lines.append(paragraph.text)
        elif code_started:
            break

    if not code_lines:
        raise SystemExit(
            f"Document validation failed: no repository code followed {heading_text}"
        )
    return code_lines


def validate_document(
    output: Path,
    root: Path,
    source_files: list[Path],
    repository_url: str,
    commit_hash: str,
) -> None:
    if output.stat().st_size < 10_000:
        raise SystemExit("Document validation failed: output file is unexpectedly small.")

    document = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required_text = [
        COURSE,
        COURSE_CODE,
        "Hands-on Assignment 1:",
        "Construct Your First Flutter App using Dart",
        STUDENT,
        "University of the Cumberlands",
        repository_url,
        commit_hash,
        "1. Application Overview",
        "2. Application Output",
        "3. Implementation Design",
        "4. Dart Source Code",
        "5. Android Manifest",
        "6. Testing and Validation",
        "7. Coding Practices",
        "8. GitHub Repository",
        "9. Development Tools and AI Assistance",
        "Figure 1. Measures Converter running on the Android emulator.",
    ]
    missing_text = [item for item in required_text if item not in body_text]
    if missing_text:
        raise SystemExit(
            "Document validation failed: missing required text: "
            + ", ".join(missing_text)
        )

    header_text = "\n".join(
        paragraph.text for paragraph in document.sections[0].header.paragraphs
    )
    footer_text = "\n".join(
        paragraph.text for paragraph in document.sections[0].footer.paragraphs
    )
    if COURSE_CODE not in header_text or "Software Engineering" not in header_text:
        raise SystemExit("Document validation failed: UC course header is incomplete.")
    if STUDENT not in footer_text or "University of the Cumberlands" not in footer_text:
        raise SystemExit("Document validation failed: UC footer is incomplete.")

    if len(document.inline_shapes) < 1:
        raise SystemExit("Document validation failed: screenshot is not embedded.")

    for source_file in source_files:
        relative_path = source_file.relative_to(root).as_posix()
        expected_lines = source_file.read_text(encoding="utf-8").splitlines()
        if repository_code_lines(document, relative_path) != expected_lines:
            raise SystemExit(
                f"Document validation failed: source does not match {relative_path}."
            )

    manifest_lines = (root / MANIFEST_PATH).read_text(encoding="utf-8").splitlines()
    if repository_code_lines(document, "5. Android Manifest") != manifest_lines:
        raise SystemExit("Document validation failed: AndroidManifest.xml does not match.")

    if not document.tables:
        raise SystemExit("Document validation failed: validation table is missing.")
    table_headers = [cell.text for cell in document.tables[0].rows[0].cells]
    if table_headers != ["Validation", "Command / Check", "Result"]:
        raise SystemExit("Document validation failed: validation table headers differ.")


def main() -> None:
    root = repository_root()
    source_files = require_inputs(root)
    remote_url = run_git(root, "remote", "get-url", "origin")
    repository_url = browser_repository_url(remote_url)
    branch = run_git(root, "branch", "--show-current")
    commit_hash = run_git(root, "rev-parse", "HEAD")

    if not branch:
        raise SystemExit("A named Git branch is required; detached HEAD was detected.")

    output = root / OUTPUT_PATH
    output.parent.mkdir(parents=True, exist_ok=True)
    document = build_document(
        root,
        source_files,
        repository_url,
        branch,
        commit_hash,
    )
    document.save(output)

    if not output.is_file() or output.stat().st_size == 0:
        raise SystemExit(f"Document generation failed: {OUTPUT_PATH}")

    validate_document(
        output,
        root,
        source_files,
        repository_url,
        commit_hash,
    )

    print(f"Generated {OUTPUT_PATH} ({output.stat().st_size:,} bytes)")
    print("Validated required content, embedded screenshot, source, and manifest.")
    print(f"Repository: {repository_url}")
    print(f"Branch: {branch}")
    print(f"Commit: {commit_hash}")
    print("Included Dart source files:")
    for source_file in source_files:
        print(f"- {source_file.relative_to(root).as_posix()}")


if __name__ == "__main__":
    main()
